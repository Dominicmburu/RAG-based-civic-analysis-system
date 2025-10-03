from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
from dotenv import load_dotenv
import faiss
import pandas as pd
from sentence_transformers import SentenceTransformer
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_openai import ChatOpenAI
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import fitz  # PyMuPDF
import docx as docx_lib
import re
import nltk
from nltk.tokenize import sent_tokenize
import uuid
from datetime import datetime
import io
import time

load_dotenv()
nltk.download('punkt', quiet=True)
nltk.download('punkt_tab', quiet=True)

app = Flask(__name__)
CORS(app)

# Global models
embedding_model = None
llm = None
user_documents = {}  # Store user-specific indices {user_id: {index, chunks, metadata}}

UPLOAD_FOLDER = 'uploads'
SUMMARIES_FOLDER = 'summaries'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(SUMMARIES_FOLDER, exist_ok=True)

def initialize_models():
    """Initialize embedding and LLM models"""
    global embedding_model, llm
    
    print("\n" + "="*60)
    print("🚀 INITIALIZING MODELS")
    print("="*60)
    
    print("📥 Loading embedding model (all-mpnet-base-v2)...")
    start = time.time()
    embedding_model = SentenceTransformer('all-mpnet-base-v2')
    print(f"✅ Embedding model loaded in {time.time()-start:.2f}s")
    
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY not set")
    
    print("🤖 Initializing GPT-4o...")
    start = time.time()
    llm = ChatOpenAI(model_name="gpt-4o", temperature=0.3)
    print(f"✅ GPT-4o initialized in {time.time()-start:.2f}s")
    print("="*60)
    print("✅ ALL MODELS READY!")
    print("="*60 + "\n")

def extract_text_from_pdf(file_path):
    """Extract text from PDF"""
    text = ""
    with fitz.open(file_path) as pdf:
        for page in pdf:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX"""
    doc = docx_lib.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def clean_text(text):
    """Clean and normalize text"""
    text = re.sub(r'\s+', ' ', text)
    text = text.replace(""", '"').replace(""", '"').replace("'", "'").replace("'", "'")
    return text.strip()

def chunk_text(text, max_tokens=300, overlap_sentences=2):
    """Chunk text into manageable pieces"""
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = []
    current_length = 0
    
    for sentence in sentences:
        token_count = len(sentence.split())
        
        if current_length + token_count <= max_tokens:
            current_chunk.append(sentence)
            current_length += token_count
        else:
            if current_chunk:
                chunks.append(" ".join(current_chunk))
            overlap = current_chunk[-overlap_sentences:] if overlap_sentences else []
            current_chunk = overlap + [sentence]
            current_length = sum(len(s.split()) for s in current_chunk)
    
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    
    return chunks

def create_user_index(user_id, file_path, filename):
    """Create FAISS index for uploaded document"""
    print("\n" + "="*70)
    print(f"📄 PROCESSING DOCUMENT: {filename}")
    print(f"   User ID: {user_id}")
    print("="*70)
    
    # Extract text based on file type
    print("📖 Step 1/5: Extracting text from document...")
    start = time.time()
    if filename.endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif filename.endswith('.docx'):
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format")
    word_count = len(text.split())
    print(f"   ✅ Extracted {word_count:,} words in {time.time()-start:.2f}s")
    
    # Clean and chunk
    print("🧹 Step 2/5: Cleaning text...")
    start = time.time()
    text = clean_text(text)
    print(f"   ✅ Text cleaned in {time.time()-start:.2f}s")
    
    print("✂️  Step 3/5: Chunking text...")
    start = time.time()
    chunks = chunk_text(text)
    print(f"   ✅ Created {len(chunks)} chunks in {time.time()-start:.2f}s")
    
    if not chunks:
        raise ValueError("No text extracted from document")
    
    # Create embeddings
    print("🧠 Step 4/5: Creating embeddings...")
    start = time.time()
    embeddings = embedding_model.encode(
        chunks,
        show_progress_bar=False,
        convert_to_numpy=True,
        normalize_embeddings=True
    )
    print(f"   ✅ Embeddings created in {time.time()-start:.2f}s")
    
    # Create FAISS index
    print("📊 Step 5/5: Building FAISS index...")
    start = time.time()
    index = faiss.IndexFlatIP(embeddings.shape[1])
    index.add(embeddings)
    print(f"   ✅ Index built in {time.time()-start:.2f}s")
    
    # Store metadata
    metadata = {
        'filename': filename,
        'chunks': chunks,
        'upload_date': datetime.now().isoformat(),
        'total_chunks': len(chunks)
    }
    
    user_documents[user_id] = {
        'index': index,
        'metadata': metadata
    }
    
    print("="*70)
    print(f"✅ DOCUMENT PROCESSING COMPLETE!")
    print(f"   Total chunks: {len(chunks)}")
    print(f"   Total words: {word_count:,}")
    print("="*70 + "\n")
    
    return metadata

def semantic_search(user_id, query, k=5):
    """Search user's document"""
    if user_id not in user_documents:
        return []
    
    doc_data = user_documents[user_id]
    index = doc_data['index']
    chunks = doc_data['metadata']['chunks']
    
    query_embedding = embedding_model.encode(
        [query],
        convert_to_numpy=True,
        normalize_embeddings=True
    )
    
    distances, indices = index.search(query_embedding, k)
    
    results = []
    for score, idx in zip(distances[0], indices[0]):
        if idx < len(chunks):
            results.append({
                'chunk_text': chunks[idx],
                'score': float(score),
                'chunk_index': int(idx)
            })
    
    return results

def create_docx_summary(title, summary_text):
    """Create formatted DOCX document"""
    doc = Document()
    
    # Title
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph("Generated by Constitutional Q&A System")
    doc.add_paragraph()
    
    # Content
    for line in summary_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith('**') and line.endswith('**'):
            # Section heading
            heading_text = line.strip('*')
            doc.add_heading(heading_text, level=2)
        else:
            doc.add_paragraph(line)
    
    return doc

# API Routes

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy", "message": "Constitution Q&A API running"})

@app.route('/upload', methods=['POST'])
def upload_document():
    """Upload and process constitutional document"""
    print("\n" + "🔵"*35)
    print("📤 NEW UPLOAD REQUEST RECEIVED")
    print("🔵"*35)
    
    if 'file' not in request.files:
        print("❌ Error: No file provided")
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    if file.filename == '':
        print("❌ Error: No file selected")
        return jsonify({"error": "No file selected"}), 400
    
    # Generate user ID (in production, use authenticated user ID)
    user_id = request.form.get('user_id', str(uuid.uuid4()))
    
    # Save file
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, f"{user_id}_{filename}")
    
    print(f"💾 Saving file: {filename}")
    file.save(filepath)
    print(f"✅ File saved to: {filepath}")
    
    try:
        metadata = create_user_index(user_id, filepath, filename)
        return jsonify({
            "user_id": user_id,
            "message": "Document processed successfully",
            "metadata": metadata
        })
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}\n")
        return jsonify({"error": str(e)}), 500

@app.route('/ask', methods=['POST'])
def ask_question():
    """Answer questions about the document"""
    data = request.json
    user_id = data.get('user_id')
    question = data.get('question')
    
    if not user_id or not question:
        return jsonify({"error": "user_id and question required"}), 400
    
    if user_id not in user_documents:
        return jsonify({"error": "Document not found. Please upload first."}), 404
    
    print(f"\n❓ Question received: {question}")
    
    # Search relevant chunks
    results = semantic_search(user_id, question, k=5)
    
    if not results:
        return jsonify({"answer": "I couldn't find relevant information in the document."})
    
    # Prepare context
    context = "\n\n".join([r['chunk_text'] for r in results])
    
    # Create QA prompt
    qa_prompt = PromptTemplate(
        input_variables=["context", "question"],
        template="""You are an expert legal assistant helping users understand constitutional documents.

Context from the document:
{context}

Question: {question}

Provide a clear, accurate answer based ONLY on the context provided. If the context doesn't contain enough information to fully answer the question, say so. Include relevant quotes from the document when appropriate.

Answer:"""
    )
    
    chain = LLMChain(llm=llm, prompt=qa_prompt)
    answer = chain.run({"context": context, "question": question})
    
    print(f"✅ Answer generated\n")
    
    return jsonify({
        "question": question,
        "answer": answer,
        "sources": [{"text": r['chunk_text'][:200] + "...", "relevance": r['score']} for r in results[:3]]
    })

@app.route('/summary', methods=['POST'])
def generate_summary():
    """Generate full document summary"""
    data = request.json
    user_id = data.get('user_id')
    
    if not user_id:
        return jsonify({"error": "user_id required"}), 400
    
    if user_id not in user_documents:
        return jsonify({"error": "Document not found"}), 404
    
    print(f"\n📝 Generating summary for user: {user_id}")
    
    doc_data = user_documents[user_id]
    chunks = doc_data['metadata']['chunks']
    filename = doc_data['metadata']['filename']
    
    # Combine all chunks (or sample if too large)
    max_chunks = 50  # Limit to avoid token overflow
    selected_chunks = chunks[:max_chunks] if len(chunks) > max_chunks else chunks
    full_text = "\n\n".join(selected_chunks)
    
    # Summary prompt
    summary_prompt = PromptTemplate(
        input_variables=["document", "filename"],
        template="""You are a legal expert summarizing a constitutional document.

Document: {filename}

Content:
{document}

Create a comprehensive summary with the following structure:

**EXECUTIVE SUMMARY**
A 2-3 paragraph overview of the document's purpose and key provisions.

**KEY PROVISIONS**
Organize main provisions by theme (e.g., Rights and Freedoms, Government Structure, Judiciary, etc.)

**NOTABLE FEATURES**
Highlight unique or significant aspects of this constitution.

**CONCLUSION**
Brief statement on the document's significance.

Keep the summary clear, objective, and accessible to general readers while maintaining legal accuracy.

Summary:"""
    )
    
    chain = LLMChain(llm=llm, prompt=summary_prompt)
    summary = chain.run({"document": full_text, "filename": filename})
    
    print(f"✅ Summary generated\n")
    
    return jsonify({
        "summary": summary,
        "filename": filename,
        "chunks_analyzed": len(selected_chunks),
        "total_chunks": len(chunks)
    })

@app.route('/download_summary', methods=['POST'])
def download_summary():
    """Download summary as DOCX"""
    data = request.json
    user_id = data.get('user_id')
    summary_text = data.get('summary')
    
    if not user_id or not summary_text:
        return jsonify({"error": "user_id and summary required"}), 400
    
    if user_id not in user_documents:
        return jsonify({"error": "Document not found"}), 404
    
    filename = user_documents[user_id]['metadata']['filename']
    title = f"Summary of {filename}"
    
    # Create DOCX
    doc = create_docx_summary(title, summary_text)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Generate download filename
    download_name = f"summary_{filename.rsplit('.', 1)[0]}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=download_name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/documents/<user_id>', methods=['GET'])
def get_user_document_info(user_id):
    """Get information about user's uploaded document"""
    if user_id not in user_documents:
        return jsonify({"error": "Document not found"}), 404
    
    metadata = user_documents[user_id]['metadata']
    return jsonify(metadata)

@app.route('/documents/<user_id>', methods=['DELETE'])
def delete_user_document(user_id):
    """Delete user's document and index"""
    if user_id not in user_documents:
        return jsonify({"error": "Document not found"}), 404
    
    print(f"\n🗑️  Deleting document for user: {user_id}")
    
    del user_documents[user_id]
    
    # Clean up file
    for file in os.listdir(UPLOAD_FOLDER):
        if file.startswith(user_id):
            os.remove(os.path.join(UPLOAD_FOLDER, file))
    
    print(f"✅ Document deleted\n")
    
    return jsonify({"message": "Document deleted successfully"})

# Initialize models before first request
with app.app_context():
    initialize_models()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)